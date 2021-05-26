from docx import Document


def get_paragraph_text(path, n):
    """
    获取指定段落的文本
    :param path: word文件路径
    :param n: 第几段落，从0开始计数
    :return: word文本
    """
    document = Document(path)
    all_paragraphs = len(document.paragraphs)
    if all_paragraphs > n:
        paragraph_text = document.paragraphs[n].text
        return paragraph_text
    else:
        raise IndexError('paragraph index (%s) out of range, in total %s' % (n, all_paragraphs))


def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path)
    all_paragraphs = document.paragraphs
    paragraphs_text = []
    for paragraph in all_paragraphs:
        paragraphs_text.append(paragraph.text)
    return paragraphs_text


def get_all_tables_text(path):
    """
    获取word中所有表格的文本
    :param path: word路径
    :return: list类型的二维数组
        如：[['年龄', '排序'], ['23', '00',], ...]
    """
    document = Document(path)
    all_tables = document.tables
    text_list = []
    for table in all_tables:
        for row in table.rows:
            text = []
            for cell in row.cells:
                text.append(cell.text)
            text_list.append(text)
    return text_list


def get_table_text(path, n=0):
    """
    获取word中的第n个表格的文本
    :param path: word路径
    :param n: 第几个表格，从0开始计算
    :return: list类型的二维数组
        如：[['年龄', '排序'], ['23', '00',], ...]
    """
    document = Document(path)
    all_tables = len(document.tables)
    if all_tables > n:
        table = document.tables[n]
        text_list = []
        for row in table.rows:
            text = []
            for cell in row.cells:
                text.append(cell.text)
            text_list.append(text)
        return text_list
    else:
        raise IndexError('table index (%s) out of range, in total %s' % (n, all_tables))


def get_cell_text(path, n=0, row=0, col=0):
    """
    获取某个表格的某个单元格的值
    :param path: word路径
    :param n: 第几个表格，从0开始计算
    :param row: 第几行，从0开始计算
    :param col: 第几列，从0开始计算
    :return: 单元格的值，str类型
    """
    document = Document(path)
    all_tables = len(document.tables)
    if all_tables > n:
        rows = len(document.tables[n].rows)
        cols = len(document.tables[n].columns)
        if rows > row and cols > col:
            tab = document.tables[n].rows[row].cells[col]
            return tab.text
        else:
            raise IndexError('cell index out of range, %s;%s' % (row, col))
    else:
        raise IndexError('table index (%s) out of range, in toatl %s' % (n, all_tables))