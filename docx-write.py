from docx import Document


def write_docx():
    doc = Document()
    doc.add_heading("Python是什么？")

    doc.add_heading("这是一级标题", level=1)
    doc.add_heading("这是二级标题", level=2)
    doc.add_heading("这是三级标题", level=3)
    doc.add_paragraph("这是一个副标题", "Subtitle")
    doc.add_paragraph("这是正文")

    doc.add_table(rows=5, cols=5)

    doc.save("res//python_word.docx")


if __name__ == '__main__':
    write_docx()
