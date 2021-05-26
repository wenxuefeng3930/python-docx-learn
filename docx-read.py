import docx


def read_docx():
    doc = docx.Document(r"res\python_word.docx")
    print("檔案內含段落數：", len(doc.paragraphs), "\n")

    test_list = []
    for text in doc.paragraphs:
        test_list.append(text)

    for item in test_list:
        print(item.text)


if __name__ == '__main__':
    read_docx()
